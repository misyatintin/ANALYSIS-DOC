"""OpenRouter API service - Full Featured Document Analysis"""
import os
import base64
import httpx
import json
import io
from dotenv import load_dotenv

load_dotenv()

OPENROUTER_API_URL = "https://openrouter.ai/api/v1/chat/completions"
API_KEY = os.getenv("OPENROUTER_API_KEY")

PDF_MODEL = "anthropic/claude-sonnet-4"
TEXT_MODEL = "openai/gpt-4o"

def extract_text_from_docx(file_data: bytes) -> str:
    """Extract text from DOCX file"""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_data))
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_parts.append(" | ".join(row_text))
        return "\n\n".join(text_parts)
    except Exception as e:
        print(f"Error extracting DOCX text: {e}")
        return ""

def get_mime_type(file_type: str) -> str:
    mime_types = {
        "pdf": "application/pdf",
        "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "doc": "application/msword",
        "png": "image/png",
        "jpg": "image/jpeg",
        "jpeg": "image/jpeg",
        "gif": "image/gif",
        "webp": "image/webp",
        "txt": "text/plain",
        "csv": "text/csv",
        "md": "text/markdown",
    }
    return mime_types.get(file_type.lower(), "application/octet-stream")

def build_file_content(file_data: bytes, filename: str, file_type: str):
    """Build content array for API request"""
    
    # Handle DOCX - extract text since Claude doesn't support DOCX files directly
    if file_type.lower() in ["docx", "doc"]:
        text_content = extract_text_from_docx(file_data)
        if text_content:
            return [{"type": "text", "text": f"Analyze this document:\n\nFilename: {filename}\n\nContent:\n{text_content[:50000]}"}], TEXT_MODEL
        else:
            return [{"type": "text", "text": f"Unable to extract text from: {filename}"}], TEXT_MODEL
    
    base64_data = base64.b64encode(file_data).decode("utf-8")
    mime_type = get_mime_type(file_type)
    
    if file_type.lower() in ["png", "jpg", "jpeg", "gif", "webp"]:
        return [
            {"type": "text", "text": f"Analyze this image: {filename}"},
            {"type": "image_url", "image_url": {"url": f"data:{mime_type};base64,{base64_data}"}}
        ], TEXT_MODEL
    elif file_type.lower() == "pdf":
        return [
            {"type": "text", "text": f"Analyze this document: {filename}"},
            {"type": "file", "file": {"filename": filename, "file_data": f"data:{mime_type};base64,{base64_data}"}}
        ], PDF_MODEL
    else:
        try:
            text_content = file_data.decode("utf-8")
            return [{"type": "text", "text": f"Analyze this document:\n\nFilename: {filename}\n\nContent:\n{text_content}"}], TEXT_MODEL
        except:
            return [{"type": "text", "text": f"Unable to read file: {filename}"}], TEXT_MODEL

async def call_openrouter(system_prompt: str, content: list, model: str, retries: int = 2) -> dict:
    """Make API call to OpenRouter with retry logic"""
    last_error = None
    
    for attempt in range(retries + 1):
        try:
            async with httpx.AsyncClient(timeout=180.0) as client:
                response = await client.post(
                    OPENROUTER_API_URL,
                    headers={
                        "Authorization": f"Bearer {API_KEY}",
                        "Content-Type": "application/json",
                        "HTTP-Referer": "https://analysisdoc.app",
                        "X-Title": "AnalysisDoc Web",
                    },
                    json={
                        "model": model,
                        "messages": [
                            {"role": "system", "content": system_prompt},
                            {"role": "user", "content": content}
                        ],
                        "response_format": {"type": "json_object"},
                    }
                )
                
                if response.status_code != 200:
                    error_text = response.text
                    print(f"OpenRouter API error (attempt {attempt + 1}): {response.status_code} - {error_text[:500]}")
                    last_error = Exception(f"OpenRouter API error: {response.status_code}")
                    if attempt < retries:
                        continue
                    raise last_error
                
                data = response.json()
                
                # Check if we have valid response
                if not data.get("choices") or not data["choices"][0].get("message"):
                    print(f"Invalid response structure (attempt {attempt + 1}): {str(data)[:500]}")
                    last_error = Exception("Invalid API response structure")
                    if attempt < retries:
                        continue
                    raise last_error
                
                content_str = data["choices"][0]["message"]["content"]
                
                # Handle empty content
                if not content_str or content_str.strip() == "":
                    print(f"Empty content received (attempt {attempt + 1})")
                    last_error = Exception("Empty response from AI")
                    if attempt < retries:
                        continue
                    raise last_error
                
                # Try to parse JSON
                try:
                    return json.loads(content_str)
                except json.JSONDecodeError as e:
                    print(f"JSON parse error (attempt {attempt + 1}): {str(e)}")
                    print(f"Content received: {content_str[:500]}")
                    # Try to extract JSON from the response
                    import re
                    json_match = re.search(r'\{[\s\S]*\}', content_str)
                    if json_match:
                        try:
                            return json.loads(json_match.group())
                        except:
                            pass
                    last_error = Exception(f"Failed to parse AI response as JSON")
                    if attempt < retries:
                        continue
                    raise last_error
                    
        except httpx.TimeoutException:
            print(f"Timeout (attempt {attempt + 1})")
            last_error = Exception("Request timed out")
            if attempt < retries:
                continue
            raise last_error
        except Exception as e:
            if "OpenRouter" in str(e) or "Empty" in str(e) or "JSON" in str(e):
                raise
            print(f"Unexpected error (attempt {attempt + 1}): {str(e)}")
            last_error = e
            if attempt < retries:
                continue
            raise
    
    raise last_error or Exception("Unknown error")

# ============ ANALYSIS FUNCTIONS ============

async def analyze_summarize(file_data: bytes, filename: str, file_type: str) -> dict:
    """Summarize a document"""
    content, model = build_file_content(file_data, filename, file_type)
    prompt = """You are a document analysis expert. Analyze the document and return JSON:
{
    "title": "Document title",
    "summary": "Comprehensive summary paragraph",
    "key_highlights": ["highlight 1", "highlight 2", "highlight 3"],
    "key_points": [
        {"label": "Point title", "details": "Detailed explanation", "page": 1, "confidence": 0.9}
    ],
    "sections": [
        {"title": "Section name", "summary": "Section summary", "page": 1}
    ],
    "document_type": "Type (report, contract, article, etc.)",
    "language": "Detected language",
    "page_count": 0,
    "word_count": 0
}
Be thorough. Include confidence scores (0.0-1.0) for each key point."""
    return await call_openrouter(prompt, content, model)

async def analyze_pros_cons(file_data: bytes, filename: str, file_type: str) -> dict:
    """Generate pros and cons analysis"""
    content, model = build_file_content(file_data, filename, file_type)
    prompt = """Analyze the document for pros and cons. Return JSON:
{
    "title": "Document title",
    "summary": "Brief overview",
    "pros": [
        {"point": "Pro description", "importance": "high/medium/low", "citation": {"page": 1, "quote": "relevant quote max 25 words", "confidence": 0.9}}
    ],
    "cons": [
        {"point": "Con description", "importance": "high/medium/low", "citation": {"page": 1, "quote": "relevant quote max 25 words", "confidence": 0.9}}
    ],
    "overall_assessment": "Overall assessment",
    "recommendation": "Final recommendation"
}
Each pro/con MUST have a citation with page number and quote."""
    return await call_openrouter(prompt, content, model)

async def analyze_gaps_risks(file_data: bytes, filename: str, file_type: str) -> dict:
    """Identify gaps and risks"""
    content, model = build_file_content(file_data, filename, file_type)
    prompt = """Analyze the document for gaps and risks. Return JSON:
{
    "title": "Document title",
    "summary": "Brief overview",
    "gaps": [
        {"description": "Gap description", "severity": "high/medium/low", "recommendation": "How to address", "citation": {"page": 1, "quote": "relevant quote", "confidence": 0.8}}
    ],
    "risks": [
        {"description": "Risk description", "severity": "high/medium/low", "impact": "Potential impact", "mitigation": "Mitigation strategy", "citation": {"page": 1, "quote": "relevant quote", "confidence": 0.8}}
    ],
    "completeness_score": 75,
    "missing_sections": ["List of missing sections"],
    "improvement_priority": ["Ordered list of improvements"]
}"""
    return await call_openrouter(prompt, content, model)

async def analyze_upgrade_suggestions(file_data: bytes, filename: str, file_type: str) -> dict:
    """Generate upgrade suggestions"""
    content, model = build_file_content(file_data, filename, file_type)
    prompt = """Analyze the document and suggest improvements. Return JSON:
{
    "title": "Document title",
    "current_quality_score": 70,
    "suggestions": [
        {
            "suggestion": "Actionable improvement",
            "priority": "high/medium/low",
            "effort": "low/medium/high",
            "impact": "Expected impact",
            "source_gap": "Related gap or issue",
            "citation": {"page": 1, "quote": "relevant quote", "confidence": 0.8}
        }
    ],
    "quick_wins": ["Easy improvements"],
    "major_improvements": ["Significant changes needed"],
    "potential_quality_score": 90
}"""
    return await call_openrouter(prompt, content, model)

async def analyze_qa(file_data: bytes, filename: str, file_type: str, question: str) -> dict:
    """Answer a question about the document"""
    content, model = build_file_content(file_data, filename, file_type)
    content[0]["text"] = f"Document: {filename}\n\nQuestion: {question}"
    prompt = """Answer the question based on the document. Return JSON:
{
    "question": "The question asked",
    "answer": "Detailed answer based on document content",
    "confidence": 0.9,
    "citations": [
        {"page": 1, "quote": "Supporting quote max 25 words", "confidence": 0.9}
    ],
    "related_topics": ["Related topics in the document"],
    "follow_up_questions": ["Suggested follow-up questions"],
    "warning": null
}
If the question cannot be answered from the document, set warning to explain why and provide best effort answer."""
    return await call_openrouter(prompt, content, model)

async def generate_chart_data(file_data: bytes, filename: str, file_type: str, chart_type: str = "bar") -> dict:
    """Extract data for chart generation"""
    content, model = build_file_content(file_data, filename, file_type)
    prompt = f"""Extract numeric data from the document suitable for a {chart_type} chart. Return JSON:
{{
    "title": "Chart title",
    "chart_type": "{chart_type}",
    "x_label": "X-axis label",
    "y_label": "Y-axis label",
    "data": [
        {{"label": "Category/Item", "value": 100}},
        {{"label": "Category/Item 2", "value": 80}}
    ],
    "source_page": 1,
    "notes": "Data source notes",
    "alternative_charts": ["Other suitable chart types"]
}}
Extract meaningful numeric data. If no numeric data found, create summary statistics."""
    return await call_openrouter(prompt, content, model)

# ============ COMPARISON FUNCTIONS ============

def build_comparison_content(doc_data: bytes, doc_name: str, doc_type: str, version_label: str):
    """Build content for a single document in comparison"""
    if doc_type.lower() in ["docx", "doc"]:
        text_content = extract_text_from_docx(doc_data)
        if text_content:
            return {"type": "text", "text": f"\n\n--- {version_label}: {doc_name} ---\n{text_content[:30000]}"}
        return {"type": "text", "text": f"\n\n--- {version_label}: {doc_name} ---\n[Unable to extract text]"}
    elif doc_type.lower() == "pdf":
        base64_data = base64.b64encode(doc_data).decode("utf-8")
        mime_type = get_mime_type(doc_type)
        return {"type": "file", "file": {"filename": f"{version_label}: {doc_name}", "file_data": f"data:{mime_type};base64,{base64_data}"}}
    else:
        try:
            text_content = doc_data.decode("utf-8")
            return {"type": "text", "text": f"\n\n--- {version_label}: {doc_name} ---\n{text_content[:30000]}"}
        except:
            return {"type": "text", "text": f"\n\n--- {version_label}: {doc_name} ---\n[Unable to read]"}

async def compare_two_documents(doc1_data: bytes, doc1_name: str, doc1_type: str,
                                 doc2_data: bytes, doc2_name: str, doc2_type: str) -> dict:
    """Compare two documents with detailed table format"""
    content = [{"type": "text", "text": f"Compare these two documents in detail:\nDocument 1: {doc1_name}\nDocument 2: {doc2_name}"}]
    
    # Add documents based on their type
    content.append(build_comparison_content(doc1_data, doc1_name, doc1_type, "Document 1"))
    content.append(build_comparison_content(doc2_data, doc2_name, doc2_type, "Document 2"))
    
    # Use PDF model only if we have PDF files, otherwise use text model
    has_pdf = doc1_type.lower() == "pdf" or doc2_type.lower() == "pdf"
    model = PDF_MODEL if has_pdf else TEXT_MODEL
    
    prompt = """Compare the two documents thoroughly with detailed analysis. Return JSON:
{
    "document1": {"name": "Doc 1 name", "summary": "Detailed summary of document 1", "key_points": ["point1", "point2"]},
    "document2": {"name": "Doc 2 name", "summary": "Detailed summary of document 2", "key_points": ["point1", "point2"]},
    "comparison_table": [
        {
            "aspect": "Aspect name (e.g., Scope, Coverage, Data Quality, Completeness)",
            "document1_value": "What document 1 has for this aspect",
            "document2_value": "What document 2 has for this aspect",
            "difference": "Key difference",
            "better": "document1/document2/equal"
        }
    ],
    "detailed_differences": [
        {
            "category": "Category (Content, Structure, Data, Coverage, etc.)",
            "type": "added/removed/modified/different",
            "severity": "high/medium/low",
            "description": "Detailed description of the difference",
            "document1_detail": "What document 1 says/has",
            "document2_detail": "What document 2 says/has",
            "impact": "Impact of this difference"
        }
    ],
    "similarity_score": 75,
    "strengths_doc1": ["Strengths of document 1"],
    "strengths_doc2": ["Strengths of document 2"],
    "weaknesses_doc1": ["Weaknesses of document 1"],
    "weaknesses_doc2": ["Weaknesses of document 2"],
    "best_version": "document1/document2",
    "best_version_reason": "Detailed explanation of why this document is better",
    "recommendation": "Final recommendation for which document to use and why"
}"""
    return await call_openrouter(prompt, content, model)

async def compare_multiple_documents(documents: list) -> dict:
    """Compare multiple documents with detailed table format"""
    content = [{"type": "text", "text": f"Compare these {len(documents)} documents in detail:"}]
    
    has_pdf = False
    for i, doc in enumerate(documents):
        content.append(build_comparison_content(doc["data"], doc["name"], doc["type"], f"Document {i+1}"))
        if doc["type"].lower() == "pdf":
            has_pdf = True
    
    model = PDF_MODEL if has_pdf else TEXT_MODEL
    
    prompt = """Compare all documents thoroughly with detailed analysis. Return JSON:
{
    "documents": [
        {"id": 1, "name": "filename", "summary": "Detailed summary", "key_points": ["point1", "point2"], "quality_score": 80}
    ],
    "comparison_table": [
        {
            "aspect": "Aspect name (Scope, Coverage, Data Quality, Completeness, etc.)",
            "values": {"doc1": "value", "doc2": "value", "doc3": "value"},
            "best": "doc1/doc2/doc3"
        }
    ],
    "detailed_differences": [
        {
            "category": "Category",
            "description": "What differs between documents",
            "by_document": {"doc1": "detail", "doc2": "detail", "doc3": "detail"},
            "severity": "high/medium/low"
        }
    ],
    "strengths_by_document": {"doc1": ["strengths"], "doc2": ["strengths"]},
    "weaknesses_by_document": {"doc1": ["weaknesses"], "doc2": ["weaknesses"]},
    "ranking": [
        {"rank": 1, "document": "doc name", "score": 85, "reason": "Why ranked here"}
    ],
    "best_candidate": {
        "name": "filename",
        "reason": "Detailed explanation of why this is the best",
        "key_advantages": ["advantages over others"]
    },
    "recommendation": "Final detailed recommendation"
}"""
    return await call_openrouter(prompt, content, model)

# ============ DECISION MATRIX ============

async def build_decision_matrix(documents: list, criteria: list) -> dict:
    """Build a detailed decision matrix comparing documents against criteria"""
    content = [{"type": "text", "text": f"Evaluate these {len(documents)} documents/options against the following criteria:\n\nCriteria:\n"}]
    
    for c in criteria:
        content[0]["text"] += f"- {c['name']} (weight: {c['weight']}): {c.get('description', '')}\n"
    
    content[0]["text"] += "\nDocuments to evaluate:"
    
    has_pdf = False
    for i, doc in enumerate(documents):
        content.append(build_comparison_content(doc["data"], doc["name"], doc["type"], f"Document {i+1}"))
        if doc["type"].lower() == "pdf":
            has_pdf = True
    
    model = PDF_MODEL if has_pdf else TEXT_MODEL
    
    criteria_json = json.dumps(criteria)
    prompt = f"""Evaluate each document against ALL criteria with detailed scoring. Criteria: {criteria_json}

Return JSON with DETAILED analysis:
{{
    "criteria": [
        {{"name": "Criterion name", "weight": 0.3, "description": "What this measures"}}
    ],
    "options": [
        {{
            "option_id": 1,
            "name": "Document name",
            "summary": "Brief summary of this document",
            "scores": [
                {{
                    "criterion": "Criterion name",
                    "score": 8,
                    "max_score": 10,
                    "weighted_score": 2.4,
                    "reason": "Detailed explanation of why this score was given",
                    "evidence": "Specific evidence from the document supporting this score"
                }}
            ],
            "total_weighted_score": 7.5,
            "strengths": ["Detailed strength 1", "Detailed strength 2"],
            "weaknesses": ["Detailed weakness 1", "Detailed weakness 2"],
            "key_findings": ["Important finding 1", "Important finding 2"]
        }}
    ],
    "comparison_by_criterion": [
        {{
            "criterion": "Criterion name",
            "weight": 0.3,
            "scores_by_document": [
                {{"document": "Doc name", "score": 8, "reason": "Why"}}
            ],
            "best_performer": "Document name",
            "analysis": "Detailed analysis of how documents compare on this criterion"
        }}
    ],
    "ranking": [
        {{
            "rank": 1,
            "document": "Document name",
            "total_score": 8.2,
            "percentage": 82,
            "summary": "Why this document ranks here"
        }}
    ],
    "winner": {{
        "name": "Winner document name",
        "total_score": 8.2,
        "percentage": 82,
        "reason": "Detailed explanation of why this document is the best choice",
        "key_advantages": ["Advantage 1", "Advantage 2"],
        "considerations": "Any caveats or considerations when choosing this option"
    }},
    "recommendation": "Final detailed recommendation explaining the decision and any context needed"
}}

Score each document 0-10 per criterion. Calculate weighted totals (score * weight). Be thorough and detailed in your analysis."""
    return await call_openrouter(prompt, content, model)

# ============ REPORT GENERATION ============

async def generate_report(file_data: bytes, filename: str, file_type: str) -> dict:
    """Generate a comprehensive analysis report"""
    content, model = build_file_content(file_data, filename, file_type)
    prompt = """Generate a comprehensive analysis report. Return JSON:
{
    "title": "Report title",
    "executive_summary": "Executive summary paragraph",
    "document_overview": {
        "type": "Document type",
        "purpose": "Document purpose",
        "audience": "Target audience",
        "date": "Document date if found"
    },
    "key_findings": [
        {
            "finding": "Key finding",
            "importance": "high/medium/low",
            "citation": {"page": 1, "quote": "Supporting quote"}
        }
    ],
    "analysis_sections": [
        {
            "title": "Section title",
            "content": "Section analysis",
            "key_points": ["Point 1", "Point 2"]
        }
    ],
    "recommendations": [
        {"recommendation": "Action item", "priority": "high/medium/low", "rationale": "Why"}
    ],
    "conclusion": "Concluding paragraph",
    "appendix": {
        "methodology": "How analysis was conducted",
        "limitations": ["Analysis limitations"],
        "data_sources": ["Source references"]
    }
}"""
    return await call_openrouter(prompt, content, model)

async def generate_slides(file_data: bytes, filename: str, file_type: str) -> dict:
    """Generate presentation slides outline"""
    content, model = build_file_content(file_data, filename, file_type)
    prompt = """Create a presentation slides outline. Return JSON:
{
    "title": "Presentation title",
    "subtitle": "Subtitle",
    "slides": [
        {
            "slide_number": 1,
            "title": "Slide title",
            "type": "title/content/chart/summary",
            "bullets": ["Bullet point 1", "Bullet point 2"],
            "speaker_notes": "Notes for presenter",
            "chart_suggestion": null
        }
    ],
    "total_slides": 10,
    "estimated_duration": "15 minutes",
    "key_messages": ["Main takeaways"],
    "visual_suggestions": ["Suggested visuals/graphics"]
}
Create 8-12 slides covering the document comprehensively."""
    return await call_openrouter(prompt, content, model)

# ============ MAIN ANALYZE FUNCTION ============

async def analyze_document(file_data: bytes, filename: str, file_type: str, analysis_type: str = "summarize", **kwargs) -> dict:
    """Main entry point for document analysis"""
    if analysis_type == "summarize":
        return await analyze_summarize(file_data, filename, file_type)
    elif analysis_type == "pros_cons":
        return await analyze_pros_cons(file_data, filename, file_type)
    elif analysis_type == "gaps_risks":
        return await analyze_gaps_risks(file_data, filename, file_type)
    elif analysis_type == "upgrade":
        return await analyze_upgrade_suggestions(file_data, filename, file_type)
    elif analysis_type == "qa":
        question = kwargs.get("question", "What is this document about?")
        return await analyze_qa(file_data, filename, file_type, question)
    elif analysis_type == "chart":
        chart_type = kwargs.get("chart_type", "bar")
        return await generate_chart_data(file_data, filename, file_type, chart_type)
    elif analysis_type == "report":
        return await generate_report(file_data, filename, file_type)
    elif analysis_type == "slides":
        return await generate_slides(file_data, filename, file_type)
    else:
        return await analyze_summarize(file_data, filename, file_type)


# ============ SMART SUGGESTIONS ============

async def get_analysis_suggestions(file_data: bytes, filename: str, file_type: str) -> dict:
    """Analyze document ONCE and get all suggestions - saves API credits"""
    content, model = build_file_content(file_data, filename, file_type)
    prompt = """Analyze this document comprehensively and provide suggestions for all possible analyses.

Return JSON:
{
    "document_summary": "2-3 sentence summary of the document",
    "document_type": "Type (report, proposal, contract, research, policy, financial, technical, etc.)",
    "key_topics": ["main topic 1", "main topic 2", "main topic 3"],
    "has_numeric_data": true/false,
    "has_comparative_content": true/false,
    "analysis_suggestions": [
        {
            "type": "summarize",
            "relevance": 0.95,
            "reason": "Why this analysis fits",
            "output_preview": "What user will get"
        },
        {
            "type": "pros_cons",
            "relevance": 0.8,
            "reason": "Why this analysis fits",
            "output_preview": "What user will get"
        },
        {
            "type": "gaps_risks",
            "relevance": 0.7,
            "reason": "Why this analysis fits",
            "output_preview": "What user will get"
        },
        {
            "type": "upgrade",
            "relevance": 0.6,
            "reason": "Why this analysis fits",
            "output_preview": "What user will get"
        },
        {
            "type": "report",
            "relevance": 0.85,
            "reason": "Why this analysis fits",
            "output_preview": "What user will get"
        },
        {
            "type": "slides",
            "relevance": 0.75,
            "reason": "Why this analysis fits",
            "output_preview": "What user will get"
        }
    ],
    "chart_suggestions": [
        {
            "type": "bar",
            "relevance": 0.9,
            "reason": "Why this chart fits",
            "data_description": "What data to visualize"
        }
    ],
    "suggested_questions": [
        "Good question to ask about this document",
        "Another relevant question"
    ],
    "compare_suggestions": {
        "good_to_compare_with": ["Types of documents this would compare well with"],
        "comparison_criteria": ["Suggested criteria for comparison"]
    },
    "decision_matrix_suggestions": {
        "suitable_for_matrix": true/false,
        "suggested_criteria": [
            {"name": "Criterion 1", "weight": 0.3, "description": "What to evaluate"}
        ]
    },
    "user_intent_keywords": {
        "financial": ["budget", "cost", "revenue", "profit"],
        "compliance": ["regulation", "policy", "requirement"],
        "technical": ["specification", "architecture", "implementation"],
        "strategic": ["goal", "objective", "plan", "strategy"]
    }
}

Provide ALL analysis types with relevance scores (0.0-1.0). Be thorough but concise."""
    return await call_openrouter(prompt, content, model)
